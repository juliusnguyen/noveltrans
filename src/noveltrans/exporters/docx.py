"""DOCX exporter via python-docx — title page + Heading 1 per chapter."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from noveltrans.errors import ExportError
from noveltrans.exporters.base import Exporter, meta_text
from noveltrans.models import Chapter, NovelMeta


class DocxExporter(Exporter):
    name = "docx"
    display_name = "Word (.docx)"
    extension = ".docx"

    def export(
        self,
        meta: NovelMeta,
        chapters: list[Chapter],
        out_path: Path,
        use_translation: bool = True,
        number_chapters: bool = False,
    ) -> Path:
        included, skipped = self.split_available(chapters, use_translation, number_chapters)
        if not included:
            raise ExportError("Không có chương nào để xuất (chưa tải/dịch chương nào).")

        title, description = meta_text(meta, use_translation)
        document = Document()
        document.core_properties.title = title
        document.core_properties.author = meta.author

        # --- title page
        title_par = document.add_paragraph()
        title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_par.add_run(title)
        run.bold = True
        run.font.size = None  # keep default style; Word scales Heading nicely
        title_par.style = document.styles["Title"]
        if use_translation and meta.translated_title:
            original_par = document.add_paragraph(f"Tên gốc: {meta.title}")
            original_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if meta.author:
            author_par = document.add_paragraph(meta.author)
            author_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if description:
            document.add_paragraph(description)
        document.add_paragraph(f"Nguồn: {meta.url} — xuất bởi NovelTrans")
        if skipped:
            document.add_paragraph(f"Bỏ qua {len(skipped)} chương chưa có nội dung.")
        document.add_page_break()

        # --- chapters
        for i, (_chapter, text) in enumerate(included):
            document.add_heading(text.title, level=1)
            for paragraph in text.body.split("\n\n"):
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())
            if i < len(included) - 1:
                document.add_page_break()

        out_path = Path(out_path)
        document.save(str(out_path))
        return out_path
