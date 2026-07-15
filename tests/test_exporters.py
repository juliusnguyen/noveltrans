import pytest
from docx import Document
from ebooklib import epub

from noveltrans.errors import ExportError
from noveltrans.exporters import get_exporter
from noveltrans.models import Chapter, NovelMeta


@pytest.fixture
def meta() -> NovelMeta:
    return NovelMeta(
        url="https://example.com/novel/1",
        site="example",
        title="Truyện Thử Nghiệm",
        author="Tác Giả A",
        description="Mô tả ngắn.",
    )


@pytest.fixture
def chapters() -> list[Chapter]:
    return [
        Chapter(
            index=0,
            title="第一章",
            url="u0",
            content="原文一\n\n原文二",
            translated="Đoạn một\n\nĐoạn hai",
            translated_title="Chương 1",
            target_lang="vi",
            status="translated",
        ),
        Chapter(
            index=1,
            title="第二章",
            url="u1",
            content="原文三",
            translated="",  # not translated -> skipped when exporting translation
            status="downloaded",
        ),
    ]


class TestTranslatedFrontMatter:
    def test_markdown_uses_translated_meta(self, tmp_path, meta, chapters):
        meta.translated_title = "Vợ Nhỏ Được Cưng Chiều"
        meta.translated_description = "Mô tả đã dịch."
        meta.translated_lang = "vi"
        out = get_exporter("markdown").export(meta, chapters, tmp_path / "book.md")
        text = out.read_text(encoding="utf-8")
        assert text.startswith("# Vợ Nhỏ Được Cưng Chiều")
        assert "*Tên gốc: Truyện Thử Nghiệm*" in text
        assert "Mô tả đã dịch." in text
        assert "Mô tả ngắn." not in text

    def test_original_export_keeps_chinese_meta(self, tmp_path, meta, chapters):
        meta.translated_title = "Vợ Nhỏ"
        meta.translated_description = "Mô tả đã dịch."
        out = get_exporter("markdown").export(
            meta, chapters, tmp_path / "book.md", use_translation=False
        )
        text = out.read_text(encoding="utf-8")
        assert text.startswith("# Truyện Thử Nghiệm")
        assert "Mô tả ngắn." in text

    def test_docx_and_epub_use_translated_title(self, tmp_path, meta, chapters):
        meta.translated_title = "Vợ Nhỏ Được Cưng Chiều"
        meta.translated_description = "Mô tả đã dịch."
        docx_out = get_exporter("docx").export(meta, chapters, tmp_path / "book.docx")
        assert Document(str(docx_out)).core_properties.title == "Vợ Nhỏ Được Cưng Chiều"
        epub_out = get_exporter("epub").export(meta, chapters, tmp_path / "book.epub")
        book = epub.read_epub(str(epub_out))
        assert book.get_metadata("DC", "title")[0][0] == "Vợ Nhỏ Được Cưng Chiều"


class TestMarkdown:
    def test_export_translation(self, tmp_path, meta, chapters):
        out = get_exporter("markdown").export(meta, chapters, tmp_path / "book.md")
        text = out.read_text(encoding="utf-8")
        assert text.startswith("# Truyện Thử Nghiệm")
        assert "**Tác giả:** Tác Giả A" in text
        assert "## Chương 1" in text
        assert "Đoạn một" in text
        # untranslated chapter is skipped and noted
        assert "第二章" not in text.replace("Bỏ qua 1 chương chưa có nội dung: 第二章", "")
        assert "Bỏ qua 1 chương" in text

    def test_export_original(self, tmp_path, meta, chapters):
        out = get_exporter("markdown").export(
            meta, chapters, tmp_path / "book.md", use_translation=False
        )
        text = out.read_text(encoding="utf-8")
        assert "## 第一章" in text and "## 第二章" in text
        assert "原文三" in text

    def test_nothing_to_export(self, tmp_path, meta):
        empty = [Chapter(index=0, title="t", url="u")]
        with pytest.raises(ExportError):
            get_exporter("markdown").export(meta, empty, tmp_path / "book.md")


class TestDocx:
    def test_export_structure(self, tmp_path, meta, chapters):
        out = get_exporter("docx").export(meta, chapters, tmp_path / "book.docx")
        document = Document(str(out))
        texts = [p.text for p in document.paragraphs]
        headings = [
            p.text for p in document.paragraphs if p.style.name.startswith("Heading 1")
        ]
        assert "Truyện Thử Nghiệm" in texts
        assert headings == ["Chương 1"]
        assert "Đoạn một" in texts and "Đoạn hai" in texts
        assert document.core_properties.title == "Truyện Thử Nghiệm"


class TestEpub:
    def test_export_structure(self, tmp_path, meta, chapters):
        out = get_exporter("epub").export(meta, chapters, tmp_path / "book.epub")
        book = epub.read_epub(str(out))
        assert book.get_metadata("DC", "title")[0][0] == "Truyện Thử Nghiệm"
        assert book.get_metadata("DC", "language")[0][0] == "vi"
        docs = [
            i for i in book.get_items() if isinstance(i, epub.EpubHtml) and i.file_name.startswith("chap_")
        ]
        assert len(docs) == 1  # only the translated chapter
        content = docs[0].content
        body = content.decode("utf-8") if isinstance(content, bytes) else str(content)
        assert "Chương 1" in body and "Đoạn một" in body
        # working TOC
        assert len(book.toc) == 1

    def test_original_language_is_zh(self, tmp_path, meta, chapters):
        out = get_exporter("epub").export(
            meta, chapters, tmp_path / "book.epub", use_translation=False
        )
        book = epub.read_epub(str(out))
        assert book.get_metadata("DC", "language")[0][0] == "zh"


class TestChapterNumbering:
    """number_chapters prepends 'Chương N: ' to the title (N = index + 1), unless the
    title already opens with its own chapter marker (no doubled number)."""

    def _num(self, chapter, use_translation=True):
        from noveltrans.exporters.base import chapter_text

        return chapter_text(chapter, use_translation, number_chapters=True).title

    def test_prepends_number_to_plain_title(self):
        ch = Chapter(
            index=4,
            title="第五章 真名",
            url="u",
            content="原文",
            translated="Thân thể.",
            translated_title="On Jianyong gặp nạn",  # no leading number
            status="translated",
        )
        from noveltrans.exporters.base import chapter_text

        text = chapter_text(ch, use_translation=True, number_chapters=True)
        assert text.title == "Chương 5: On Jianyong gặp nạn"  # number + parsed title
        assert text.body == "Thân thể."  # body untouched

    def test_guard_skips_vietnamese_marker(self):
        ch = Chapter(index=0, title="t", url="u", content="c", translated="b",
                     translated_title="Chương 1: Trường học", status="translated")
        assert self._num(ch) == "Chương 1: Trường học"  # not "Chương 1: Chương 1: …"

    def test_guard_skips_chinese_marker(self):
        ch = Chapter(index=2, title="第三章 真名", url="u", content="原文", status="downloaded")
        assert self._num(ch, use_translation=False) == "第三章 真名"  # left as-is

    def test_guard_skips_english_marker(self):
        ch = Chapter(index=1, title="t", url="u", content="c", translated="b",
                     translated_title="Chapter 2 - The Fight", status="translated")
        assert self._num(ch) == "Chapter 2 - The Fight"

    def test_markdown_headings_are_numbered(self, tmp_path, meta):
        chs = [
            Chapter(index=0, title="第一章", url="u0", content="a", translated="Đoạn một",
                    translated_title="Mở đầu", target_lang="vi", status="translated"),
        ]
        out = get_exporter("markdown").export(
            meta, chs, tmp_path / "book.md", number_chapters=True
        )
        text = out.read_text(encoding="utf-8")
        assert "## Chương 1: Mở đầu" in text
        assert "Đoạn một" in text  # body preserved

    def test_numbering_is_gap_preserving(self, tmp_path, meta):
        # index 1 has no content → skipped; the surrounding chapters keep 1 and 3
        gapped = [
            Chapter(index=0, title="Mở đầu", url="u0", content="A", status="downloaded"),
            Chapter(index=1, title="Trống", url="u1", content="", status="pending"),
            Chapter(index=2, title="Kết", url="u2", content="C", status="downloaded"),
        ]
        out = get_exporter("markdown").export(
            meta, gapped, tmp_path / "book.md", use_translation=False, number_chapters=True
        )
        text = out.read_text(encoding="utf-8")
        assert "## Chương 1: Mở đầu" in text and "## Chương 3: Kết" in text
        assert "Chương 2" not in text  # skipped chapter is not renumbered away

    def test_epub_toc_is_numbered(self, tmp_path, meta):
        chs = [
            Chapter(index=0, title="第一章", url="u0", content="a", translated="Đoạn một",
                    translated_title="Mở đầu", target_lang="vi", status="translated"),
        ]
        out = get_exporter("epub").export(meta, chs, tmp_path / "book.epub", number_chapters=True)
        book = epub.read_epub(str(out))
        assert [item.title for item in book.toc] == ["Chương 1: Mở đầu"]

    def test_default_keeps_parsed_title(self, tmp_path, meta, chapters):
        # regression: without the flag, an original-language export keeps the real title
        out = get_exporter("markdown").export(
            meta, chapters, tmp_path / "book.md", use_translation=False
        )
        text = out.read_text(encoding="utf-8")
        assert "## 第一章" in text  # parsed title preserved
        assert "Chương 1:" not in text  # no numbering applied


class TestRegistry:
    def test_unknown(self):
        with pytest.raises(ExportError):
            get_exporter("pdf")
